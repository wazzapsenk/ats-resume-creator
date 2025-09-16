import re
import json
from typing import List, Dict, Any, Optional, Set, Tuple
from pathlib import Path
import docx
import PyPDF2
from datetime import datetime
import spacy
from collections import defaultdict

class EnhancedNLPService:
    """Enhanced NLP service with advanced text processing and skill matching"""

    def __init__(self):
        # Load spaCy model (if available)
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            print("Warning: spaCy model 'en_core_web_sm' not found. Using basic processing.")
            self.nlp = None

        # Enhanced skill taxonomy with synonyms and variations
        self.skill_taxonomy = self._load_skill_taxonomy()

        # Experience keywords for parsing
        self.experience_keywords = {
            'start_indicators': ['from', 'since', 'starting', 'began', 'joined'],
            'end_indicators': ['to', 'until', 'through', 'ended', 'left'],
            'current_indicators': ['present', 'current', 'now', 'ongoing', 'today'],
            'duration_patterns': [
                r'(\d+)\s*years?',
                r'(\d+)\s*months?',
                r'(\d+)\s*yrs?',
                r'(\d+)\s*mos?'
            ]
        }

        # Education level hierarchy
        self.education_levels = {
            'phd': ['phd', 'ph.d', 'doctorate', 'doctoral', 'doctor of philosophy'],
            'masters': ['masters', 'master', 'ms', 'm.s', 'mba', 'm.b.a', 'ma', 'm.a'],
            'bachelors': ['bachelors', 'bachelor', 'bs', 'b.s', 'ba', 'b.a', 'btech', 'b.tech'],
            'associates': ['associates', 'associate', 'as', 'a.s', 'aa', 'a.a'],
            'certificate': ['certificate', 'certification', 'diploma', 'bootcamp'],
            'high_school': ['high school', 'secondary', 'diploma', 'ged']
        }

        # Job seniority indicators
        self.seniority_keywords = {
            'entry': ['entry', 'junior', 'intern', 'trainee', 'graduate', 'associate', 'beginner'],
            'mid': ['mid', 'intermediate', 'regular', 'standard', 'developer', 'analyst'],
            'senior': ['senior', 'sr', 'lead', 'principal', 'expert', 'specialist'],
            'executive': ['director', 'manager', 'head', 'chief', 'cto', 'ceo', 'vp', 'vice president']
        }

    def _load_skill_taxonomy(self) -> Dict[str, Dict[str, List[str]]]:
        """Load comprehensive skill taxonomy with synonyms"""
        return {
            "programming_languages": {
                "primary_skills": [
                    "python", "java", "javascript", "typescript", "c++", "c#", "php",
                    "ruby", "go", "rust", "swift", "kotlin", "scala", "r", "matlab"
                ],
                "synonyms": {
                    "javascript": ["js", "ecmascript", "node.js", "nodejs"],
                    "typescript": ["ts"],
                    "c++": ["cpp", "c plus plus"],
                    "c#": ["csharp", "c sharp"],
                    "python": ["py"],
                    "java": ["jvm"],
                    "r": ["r language", "r programming"]
                }
            },
            "web_frameworks": {
                "primary_skills": [
                    "react", "angular", "vue", "django", "flask", "spring", "express",
                    "laravel", "rails", "next.js", "nuxt.js", "svelte", "ember"
                ],
                "synonyms": {
                    "react": ["reactjs", "react.js"],
                    "angular": ["angularjs", "angular.js"],
                    "vue": ["vuejs", "vue.js"],
                    "express": ["expressjs", "express.js"],
                    "next.js": ["nextjs"],
                    "nuxt.js": ["nuxtjs"]
                }
            },
            "databases": {
                "primary_skills": [
                    "mysql", "postgresql", "mongodb", "redis", "elasticsearch",
                    "sqlite", "oracle", "cassandra", "dynamodb", "firebase"
                ],
                "synonyms": {
                    "postgresql": ["postgres", "psql"],
                    "mongodb": ["mongo"],
                    "elasticsearch": ["elastic", "es"],
                    "dynamodb": ["dynamo"]
                }
            },
            "cloud_platforms": {
                "primary_skills": [
                    "aws", "azure", "gcp", "docker", "kubernetes", "terraform",
                    "jenkins", "heroku", "vercel", "digital ocean"
                ],
                "synonyms": {
                    "aws": ["amazon web services", "amazon aws"],
                    "gcp": ["google cloud", "google cloud platform"],
                    "azure": ["microsoft azure"],
                    "kubernetes": ["k8s"],
                    "digital ocean": ["digitalocean"]
                }
            },
            "tools_and_software": {
                "primary_skills": [
                    "git", "github", "gitlab", "jira", "confluence", "slack",
                    "figma", "photoshop", "illustrator", "sketch", "linux", "docker"
                ],
                "synonyms": {
                    "git": ["version control"],
                    "photoshop": ["ps", "adobe photoshop"],
                    "illustrator": ["ai", "adobe illustrator"]
                }
            },
            "soft_skills": {
                "primary_skills": [
                    "leadership", "communication", "teamwork", "problem solving",
                    "analytical thinking", "project management", "agile", "scrum"
                ],
                "synonyms": {
                    "problem solving": ["problem-solving", "troubleshooting"],
                    "project management": ["pm", "project coordination"],
                    "communication": ["verbal communication", "written communication"]
                }
            }
        }

    def enhanced_extract_text_from_file(self, file_path: str) -> Dict[str, Any]:
        """Enhanced text extraction with metadata"""
        file_path = Path(file_path)
        extension = file_path.suffix.lower()

        result = {
            "text": "",
            "metadata": {
                "file_type": extension,
                "extraction_method": "",
                "word_count": 0,
                "sections_detected": [],
                "extraction_quality": "unknown"
            }
        }

        try:
            if extension == '.pdf':
                result.update(self._enhanced_extract_from_pdf(file_path))
            elif extension in ['.docx', '.doc']:
                result.update(self._enhanced_extract_from_docx(file_path))
            elif extension == '.txt':
                result.update(self._enhanced_extract_from_txt(file_path))
            else:
                raise ValueError(f"Unsupported file format: {extension}")

            # Post-process text
            result["text"] = self._clean_extracted_text(result["text"])
            result["metadata"]["word_count"] = len(result["text"].split())
            result["metadata"]["sections_detected"] = self._detect_resume_sections(result["text"])
            result["metadata"]["extraction_quality"] = self._assess_extraction_quality(result["text"])

        except Exception as e:
            raise Exception(f"Failed to extract text from {file_path}: {str(e)}")

        return result

    def _enhanced_extract_from_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Enhanced PDF extraction with better text handling"""
        text = ""
        metadata = {"extraction_method": "PyPDF2", "pages_processed": 0}

        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata["pages_processed"] = len(pdf_reader.pages)

                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    # Clean up common PDF extraction issues
                    page_text = re.sub(r'([a-z])([A-Z])', r'\1 \2', page_text)  # Add space between camelCase
                    page_text = re.sub(r'(\w)(\n)(\w)', r'\1 \3', page_text)  # Join broken words
                    text += page_text + "\n"

        except Exception as e:
            metadata["extraction_error"] = str(e)

        return {"text": text, "metadata": metadata}

    def _enhanced_extract_from_docx(self, file_path: Path) -> Dict[str, Any]:
        """Enhanced DOCX extraction with formatting preservation"""
        text = ""
        metadata = {"extraction_method": "python-docx", "paragraphs_processed": 0}

        try:
            doc = docx.Document(file_path)
            metadata["paragraphs_processed"] = len(doc.paragraphs)

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"

            # Extract from tables if present
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"

        except Exception as e:
            metadata["extraction_error"] = str(e)

        return {"text": text, "metadata": metadata}

    def _enhanced_extract_from_txt(self, file_path: Path) -> Dict[str, Any]:
        """Enhanced TXT extraction with encoding detection"""
        text = ""
        metadata = {"extraction_method": "text", "encoding": "utf-8"}

        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                        metadata["encoding"] = encoding
                        break
                except UnicodeDecodeError:
                    continue

        except Exception as e:
            metadata["extraction_error"] = str(e)

        return {"text": text, "metadata": metadata}

    def _clean_extracted_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)  # Preserve paragraph breaks
        text = re.sub(r' {2,}', ' ', text)  # Multiple spaces to single space
        text = re.sub(r'\t+', ' ', text)  # Tabs to spaces

        # Fix common extraction artifacts
        text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)  # camelCase separation
        text = re.sub(r'(\w)([•·▪▫◦‣⁃])', r'\1 \2', text)  # Bullet points
        text = re.sub(r'([.!?])([A-Z])', r'\1 \2', text)  # Sentence separation

        return text.strip()

    def _detect_resume_sections(self, text: str) -> List[str]:
        """Detect common resume sections"""
        sections = []
        section_patterns = {
            'contact': r'(contact|phone|email|address|linkedin)',
            'summary': r'(summary|profile|objective|about|overview)',
            'experience': r'(experience|employment|work|career|professional)',
            'education': r'(education|academic|degree|university|college)',
            'skills': r'(skills|technical|competencies|expertise|technologies)',
            'projects': r'(projects|portfolio|work samples)',
            'certifications': r'(certifications|certificates|credentials)',
            'awards': r'(awards|honors|achievements|recognition)',
            'languages': r'(languages|linguistic)'
        }

        text_lower = text.lower()
        for section, pattern in section_patterns.items():
            if re.search(pattern, text_lower):
                sections.append(section)

        return sections

    def _assess_extraction_quality(self, text: str) -> str:
        """Assess the quality of text extraction"""
        if not text or len(text.strip()) < 50:
            return "poor"

        # Check for common extraction issues
        total_chars = len(text)
        if total_chars == 0:
            return "poor"

        # Calculate ratio of alphanumeric to total characters
        alphanumeric_chars = len(re.findall(r'[a-zA-Z0-9]', text))
        alphanumeric_ratio = alphanumeric_chars / total_chars

        # Check for excessive special characters (indicating extraction issues)
        special_chars = len(re.findall(r'[^\w\s.,!?()-]', text))
        special_ratio = special_chars / total_chars

        if alphanumeric_ratio > 0.7 and special_ratio < 0.1:
            return "excellent"
        elif alphanumeric_ratio > 0.5 and special_ratio < 0.2:
            return "good"
        elif alphanumeric_ratio > 0.3:
            return "fair"
        else:
            return "poor"

    def enhanced_extract_skills(self, text: str) -> Dict[str, Any]:
        """Enhanced skill extraction with confidence scoring"""
        text_lower = text.lower()
        extracted_skills = {}
        confidence_scores = {}

        for category, category_data in self.skill_taxonomy.items():
            skills_found = {}

            # Check primary skills
            for skill in category_data["primary_skills"]:
                confidence = self._calculate_skill_confidence(text_lower, skill)
                if confidence > 0:
                    skills_found[skill] = confidence

            # Check synonyms
            if "synonyms" in category_data:
                for primary_skill, synonyms in category_data["synonyms"].items():
                    if primary_skill not in skills_found:
                        max_confidence = 0
                        for synonym in synonyms:
                            confidence = self._calculate_skill_confidence(text_lower, synonym)
                            max_confidence = max(max_confidence, confidence)
                        if max_confidence > 0:
                            skills_found[primary_skill] = max_confidence

            # Filter by confidence threshold
            filtered_skills = {skill: conf for skill, conf in skills_found.items() if conf >= 0.3}

            if filtered_skills:
                extracted_skills[category] = list(filtered_skills.keys())
                confidence_scores[category] = filtered_skills

        return {
            "skills": extracted_skills,
            "confidence_scores": confidence_scores,
            "total_skills_found": sum(len(skills) for skills in extracted_skills.values())
        }

    def _calculate_skill_confidence(self, text: str, skill: str) -> float:
        """Calculate confidence score for skill detection"""
        skill_lower = skill.lower()

        # Exact word boundary matches
        exact_matches = len(re.findall(rf'\b{re.escape(skill_lower)}\b', text))
        if exact_matches > 0:
            # Higher confidence for multiple mentions
            return min(0.5 + (exact_matches * 0.2), 1.0)

        # Partial matches (lower confidence)
        if skill_lower in text:
            return 0.3

        return 0.0

    def enhanced_extract_experience(self, text: str) -> Dict[str, Any]:
        """Enhanced experience extraction with years calculation"""
        experience_data = {
            "total_years": 0,
            "positions": [],
            "companies": [],
            "industries": [],
            "seniority_level": "unknown"
        }

        # Extract years of experience
        years_patterns = [
            r'(\d+)\+?\s*years?\s*(?:of\s*)?experience',
            r'(\d+)\+?\s*yrs?\s*(?:of\s*)?experience',
            r'over\s*(\d+)\s*years?',
            r'more\s*than\s*(\d+)\s*years?'
        ]

        for pattern in years_patterns:
            matches = re.findall(pattern, text.lower())
            if matches:
                years = max([int(match) for match in matches])
                experience_data["total_years"] = max(experience_data["total_years"], years)

        # Extract job titles and companies
        job_title_patterns = [
            r'(?:^|\n)\s*([A-Z][a-zA-Z\s]+(?:Engineer|Developer|Manager|Analyst|Specialist|Coordinator|Director))',
            r'(?:Title|Position|Role):\s*([A-Z][a-zA-Z\s]+)',
        ]

        for pattern in job_title_patterns:
            matches = re.findall(pattern, text, re.MULTILINE)
            experience_data["positions"].extend([match.strip() for match in matches])

        # Determine seniority level
        text_lower = text.lower()
        seniority_score = {"entry": 0, "mid": 0, "senior": 0, "executive": 0}

        for level, keywords in self.seniority_keywords.items():
            for keyword in keywords:
                count = len(re.findall(rf'\b{keyword}\b', text_lower))
                seniority_score[level] += count

        if any(seniority_score.values()):
            experience_data["seniority_level"] = max(seniority_score.items(), key=lambda x: x[1])[0]

        return experience_data

    def enhanced_extract_education(self, text: str) -> Dict[str, Any]:
        """Enhanced education extraction with level detection"""
        education_data = {
            "highest_level": "unknown",
            "degrees": [],
            "institutions": [],
            "fields_of_study": [],
            "gpa": None
        }

        text_lower = text.lower()

        # Detect education levels
        level_scores = {level: 0 for level in self.education_levels.keys()}

        for level, keywords in self.education_levels.items():
            for keyword in keywords:
                matches = len(re.findall(rf'\b{keyword}\b', text_lower))
                level_scores[level] += matches

        # Determine highest education level
        if any(level_scores.values()):
            # Priority order for education levels
            priority_order = ['phd', 'masters', 'bachelors', 'associates', 'certificate', 'high_school']
            for level in priority_order:
                if level_scores[level] > 0:
                    education_data["highest_level"] = level
                    break

        # Extract GPA
        gpa_patterns = [
            r'gpa:?\s*(\d+\.?\d*)',
            r'(\d\.\d+)\s*/?\s*4\.0',
            r'(\d\.\d+)\s*gpa'
        ]

        for pattern in gpa_patterns:
            matches = re.findall(pattern, text_lower)
            if matches:
                try:
                    gpa = float(matches[0])
                    if 0 <= gpa <= 4.0:
                        education_data["gpa"] = gpa
                        break
                except ValueError:
                    continue

        # Extract universities/institutions
        university_patterns = [
            r'university of ([a-zA-Z\s]+)',
            r'([a-zA-Z\s]+) university',
            r'([a-zA-Z\s]+) college',
            r'([a-zA-Z\s]+) institute'
        ]

        for pattern in university_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            education_data["institutions"].extend([match.strip() for match in matches if len(match.strip()) > 2])

        return education_data

    def calculate_advanced_keyword_density(self, text: str, target_keywords: List[str]) -> Dict[str, Any]:
        """Advanced keyword density calculation with context analysis"""
        text_lower = text.lower()
        total_words = len(re.findall(r'\b\w+\b', text_lower))

        if total_words == 0:
            return {"densities": {}, "total_words": 0, "analysis": {}}

        densities = {}
        context_analysis = {}

        for keyword in target_keywords:
            keyword_lower = keyword.lower()

            # Exact matches
            exact_matches = len(re.findall(rf'\b{re.escape(keyword_lower)}\b', text_lower))

            # Contextual analysis - find surrounding words
            context_words = []
            pattern = rf'\b\w+\s+\w*{re.escape(keyword_lower)}\w*\s+\w+\b'
            contexts = re.findall(pattern, text_lower)

            for context in contexts:
                words = context.split()
                context_words.extend(words)

            # Calculate density
            density = (exact_matches / total_words) * 100
            densities[keyword] = {
                "density_percentage": round(density, 2),
                "count": exact_matches,
                "context_strength": len(set(context_words))  # Unique context words
            }

            context_analysis[keyword] = {
                "surrounding_words": list(set(context_words)),
                "appears_in_titles": bool(re.search(rf'\b{re.escape(keyword_lower)}\b',
                                                   text_lower.split('\n')[0] if text_lower else ""))
            }

        return {
            "densities": densities,
            "total_words": total_words,
            "context_analysis": context_analysis,
            "overall_keyword_coverage": len([k for k, v in densities.items() if v["count"] > 0]) / len(target_keywords) if target_keywords else 0
        }

    def assess_ats_compatibility(self, text: str, resume_data: Dict = None) -> Dict[str, Any]:
        """Comprehensive ATS compatibility assessment"""
        issues = []
        suggestions = []
        score_factors = {}

        # Text length check
        word_count = len(text.split())
        if word_count < 200:
            issues.append({
                "type": "content_length",
                "severity": "high",
                "description": "Resume appears too short for comprehensive ATS parsing",
                "recommendation": "Expand content to at least 300-500 words"
            })
            score_factors["length"] = 0.3
        elif word_count > 1000:
            issues.append({
                "type": "content_length",
                "severity": "medium",
                "description": "Resume may be too long for optimal ATS processing",
                "recommendation": "Consider condensing to 500-800 words"
            })
            score_factors["length"] = 0.7
        else:
            score_factors["length"] = 1.0

        # Contact information check
        contact_score = 0
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        phone_pattern = r'(?:\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})'

        if re.search(email_pattern, text):
            contact_score += 0.5
        else:
            issues.append({
                "type": "missing_contact",
                "severity": "high",
                "description": "Email address not detected",
                "recommendation": "Include a clear email address"
            })

        if re.search(phone_pattern, text):
            contact_score += 0.5
        else:
            issues.append({
                "type": "missing_contact",
                "severity": "medium",
                "description": "Phone number not clearly detected",
                "recommendation": "Include a clear phone number"
            })

        score_factors["contact"] = contact_score

        # Section structure check
        required_sections = ['experience', 'education', 'skills']
        detected_sections = self._detect_resume_sections(text)
        section_score = len([s for s in required_sections if s in detected_sections]) / len(required_sections)
        score_factors["structure"] = section_score

        if section_score < 1.0:
            missing_sections = [s for s in required_sections if s not in detected_sections]
            issues.append({
                "type": "missing_sections",
                "severity": "medium",
                "description": f"Missing standard resume sections: {', '.join(missing_sections)}",
                "recommendation": "Include clear section headers for Experience, Education, and Skills"
            })

        # Calculate overall ATS score
        weights = {"length": 0.2, "contact": 0.3, "structure": 0.3, "formatting": 0.2}
        score_factors["formatting"] = 0.8  # Default formatting score

        ats_score = sum(score_factors[factor] * weights[factor] for factor in weights.keys()) * 100

        return {
            "ats_score": round(ats_score, 1),
            "score_factors": score_factors,
            "issues": issues,
            "suggestions": suggestions,
            "detected_sections": detected_sections,
            "word_count": word_count
        }

# Global instance
enhanced_nlp_service = EnhancedNLPService()